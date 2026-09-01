import os
import sys

def build_docx():
    try:
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import parse_xml, OxmlElement
        from docx.oxml.ns import nsdecls, qn
    except ImportError:
        print("python-docx not installed in current interpreter")
        sys.exit(1)

    doc = docx.Document()

    # Set page margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Color Palette Constants
    COLOR_PRIMARY = RGBColor(0, 168, 150)    # #00A896 Teal
    COLOR_DARK = RGBColor(14, 56, 44)       # #0E382C Dark Mint
    COLOR_SECONDARY = RGBColor(42, 117, 113) # #2A7571 Cyan Text
    COLOR_TEXT = RGBColor(30, 41, 59)       # #1E293B Slate 800
    COLOR_MUTED = RGBColor(100, 116, 139)   # #64748B Slate 500

    # Helper function to set cell background color
    def set_cell_background(cell, fill_hex):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # Helper function to add styled heading
    def add_custom_heading(text, level=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Plus Jakarta Sans'
        run.bold = True
        
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = COLOR_PRIMARY
            # Add bottom line under H1
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="4" w:color="00C896"/></w:pBdr>')
            pPr.append(pBdr)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = COLOR_DARK
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = COLOR_SECONDARY
        return p

    # --- TITLE BLOCK ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("Wice Waste Management UI")
    run_title.font.name = 'Plus Jakarta Sans'
    run_title.font.size = Pt(26)
    run_title.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_p.add_run("Frontend Developer Assessment — Technical Submission Documentation")
    run_sub.font.name = 'Plus Jakarta Sans'
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = COLOR_MUTED
    run_sub.italic = True

    # --- METADATA TABLE ---
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    meta_data = [
        ("Project Objective:", "Pixel-accurate implementation of 3 Figma screens with fluid micro-interactions"),
        ("Figma Reference:", "https://www.figma.com/design/WAyvXxbQcf6B809aKgdmaC/Sample-app"),
        ("Technology Stack:", "React 18, Vite, Tailwind CSS, Framer Motion, Lucide Icons, Canvas Confetti"),
        ("Target Display:", "Cross-device Mobile & Tablet UI (100% Full Bleed & Safe-Area Aware)"),
        ("Build Status:", "Pure Static SPA — 0 Error Build Verification (Vite Production Output)")
    ]

    for i, (label, val) in enumerate(meta_data):
        row = meta_table.rows[i]
        
        # Label cell
        cell_lbl = row.cells[0]
        cell_lbl.width = Inches(1.8)
        set_cell_background(cell_lbl, "F0FAF9")
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_before = Pt(4)
        p_lbl.paragraph_format.space_after = Pt(4)
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.name = 'Plus Jakarta Sans'
        r_lbl.font.size = Pt(10)
        r_lbl.bold = True
        r_lbl.font.color.rgb = COLOR_DARK

        # Value cell
        cell_val = row.cells[1]
        cell_val.width = Inches(4.7)
        set_cell_background(cell_val, "FFFFFF")
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_before = Pt(4)
        p_val.paragraph_format.space_after = Pt(4)
        r_val = p_val.add_run(val)
        r_val.font.name = 'Plus Jakarta Sans'
        r_val.font.size = Pt(10)
        r_val.font.color.rgb = COLOR_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # --- 1. EXECUTIVE SUMMARY ---
    add_custom_heading("1. Executive Summary", level=1)
    
    p_exec = doc.add_paragraph()
    p_exec.paragraph_format.space_after = Pt(8)
    p_exec.paragraph_format.line_spacing = 1.15
    r_exec = p_exec.add_run(
        "This project is a high-fidelity, pixel-accurate frontend implementation of the Wice Waste Management mobile application, built directly from the provided Figma specifications. The application translates three key mobile screens into a modern, fluid React single-page web app featuring safe-area handling, spring micro-animations, interactive sheet modals, and dynamic state management."
    )
    r_exec.font.name = 'Plus Jakarta Sans'
    r_exec.font.size = Pt(10.5)
    r_exec.font.color.rgb = COLOR_TEXT

    # --- 2. REQUIRED SCREENS IMPLEMENTATION ---
    add_custom_heading("2. Detailed Screen Implementation", level=1)

    # Screen 1
    add_custom_heading("2.1 Screen 1: Welcome Screen (E / Welcome Screen)", level=2)
    p_s1 = doc.add_paragraph()
    p_s1.paragraph_format.space_after = Pt(6)
    p_s1.paragraph_format.line_spacing = 1.15
    r_s1 = p_s1.add_run(
        "• Background & Branding: Implemented using full-bleed bg.png background artwork containing the sky gradient, clouds, green hills, and wice logo illustration.\n"
        "• Typography System: Integrated Plus Jakarta Sans Google Font family for the primary headline 'Smart Waste management Made Easy'.\n"
        "• Interactive Actions: Floating 'Continue with Email' cyan button and 'Continue with Google' white button featuring official 4-color Google logo SVG.\n"
        "• Navigation Link: 'Already have an account? Sign In' text button seamlessly transitioning to the Schedule screen."
    )
    r_s1.font.name = 'Plus Jakarta Sans'
    r_s1.font.size = Pt(10)
    r_s1.font.color.rgb = COLOR_TEXT

    # Screen 2
    add_custom_heading("2.2 Screen 2: Schedule Screen (J / 04 / Pick Up Trash / Schedule / Default State)", level=2)
    p_s2 = doc.add_paragraph()
    p_s2.paragraph_format.space_after = Pt(6)
    p_s2.paragraph_format.line_spacing = 1.15
    r_s2 = p_s2.add_run(
        "• Header Navigation: Sticky header with back arrow button and centered 'Schedule' title text.\n"
        "• Service Card: White card with rounded-2xl corners, 40% enlarged official Trash Truck SVG icon, and 'Change' outline button.\n"
        "• Info Banner: Inline cyan notification row with circular warning badge (!) and text 'This method must have minimum of 5kg of waste weight' (#2A7571).\n"
        "• Date & Time Pickers: Individual rounded cards with Date Pickup (calendar search icon) and Time Pickup (clock icon) starting in default empty state.\n"
        "• Empty Address State: Rendered using address.png character illustration, 'No address added' title, 'Click the add address to continue' subtext, and 'add address' button.\n"
        "• Conditional CTA Behavior: Matching Figma default state, the bottom 'Continue to Confirmation' button remains hidden initially and smoothly slides up with Framer Motion when any selection is made."
    )
    r_s2.font.name = 'Plus Jakarta Sans'
    r_s2.font.size = Pt(10)
    r_s2.font.color.rgb = COLOR_TEXT

    # Screen 3
    add_custom_heading("2.3 Screen 3: Confirmation Screen (L / 05 / Confirmation / Default State)", level=2)
    p_s3 = doc.add_paragraph()
    p_s3.paragraph_format.space_after = Pt(6)
    p_s3.paragraph_format.line_spacing = 1.15
    r_s3 = p_s3.add_run(
        "• Centered Header: Sticky header with back button and centered 'Confirmation' title text.\n"
        "• Trash Bank Section: Clean floating detail view with MapPin icon, 'BSU – Ketupat Pandan' title, street/city address lines, Phone icon, and selected pickup date.\n"
        "• Total Weight Summary: White card displaying itemized waste categories (Plastic bottle - 3 Liter, Glass - 2 Item, Electronic - 3 Item, Oil - 2 Liter) with calculated 'Estimated total weight: 7.5 Kg'.\n"
        "• Process CTA & Celebration: Bottom teal 'Proses' button with Checkmark Circle icon trigger bringing up an animated Success Sheet Modal with realistic Canvas Confetti explosion."
    )
    r_s3.font.name = 'Plus Jakarta Sans'
    r_s3.font.size = Pt(10)
    r_s3.font.color.rgb = COLOR_TEXT

    # --- 3. ARCHITECTURE & COMPONENT SYSTEM ---
    add_custom_heading("3. Component Architecture & Design System", level=1)
    
    arch_table = doc.add_table(rows=6, cols=2)
    arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    arch_table.autofit = False

    arch_headers = ["Component Name", "Role & Description"]
    hdr_row = arch_table.rows[0]
    for j, htext in enumerate(arch_headers):
        cell = hdr_row.cells[j]
        set_cell_background(cell, "00A896")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(htext)
        run.font.name = 'Plus Jakarta Sans'
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    components_info = [
        ("Header.jsx", "Sticky navigation bar featuring centered page titles, back arrow button, and glassmorphism backdrop blur."),
        ("Button.jsx", "Reusable Framer Motion action button with spring micro-interactions on tap (scale: 0.97)."),
        ("SelectorRow.jsx", "Custom picker row for Date and Time selections with icon badges and cyan chevron indicators."),
        ("DatePickerModal / TimePickerModal / AddressModal", "Interactive bottom sheet modals for selecting dates, times, and editing location address data."),
        ("vectors.jsx", "Centralized vector hub containing user's official Trash Truck SVG, Google SVG, and address.png illustration.")
    ]

    for i, (cname, cdesc) in enumerate(components_info):
        row = arch_table.rows[i + 1]
        
        c0 = row.cells[0]
        c0.width = Inches(2.2)
        set_cell_background(c0, "F8F8FB" if i % 2 == 0 else "FFFFFF")
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(4)
        p0.paragraph_format.space_after = Pt(4)
        r0 = p0.add_run(cname)
        r0.font.name = 'Plus Jakarta Sans'
        r0.font.size = Pt(9.5)
        r0.bold = True
        r0.font.color.rgb = COLOR_DARK

        c1 = row.cells[1]
        c1.width = Inches(4.3)
        set_cell_background(c1, "F8F8FB" if i % 2 == 0 else "FFFFFF")
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after = Pt(4)
        r1 = p1.add_run(cdesc)
        r1.font.name = 'Plus Jakarta Sans'
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = COLOR_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # --- 4. KEY TECHNICAL HIGHLIGHTS ---
    add_custom_heading("4. Key Technical Highlights & Deliverables", level=1)
    
    p_tech = doc.add_paragraph()
    p_tech.paragraph_format.space_after = Pt(6)
    p_tech.paragraph_format.line_spacing = 1.15
    r_tech = p_tech.add_run(
        "1. Micro-Animations: Micro-tactile spring feedback on all button taps (whileTap={{ scale: 0.97 }}), page horizontal slide transitions, and confetti celebration effects.\n"
        "2. Cross-Device Responsiveness: Engineered to fit 100% of modern mobile screen viewports (iPhone, Android, Tablets) with safe area notch handling and clean vertical scrolling.\n"
        "3. Design Token Accuracy: Exact color matching (#00C896, #E6FAF4, #0E382C, #EBF7F5, #2A7571) and Plus Jakarta Sans font hierarchy.\n"
        "4. Asset Organization: Streamlined asset architecture stored under src/assets/images/ (bg.png, logo.png, address.png) with zero duplicate overhead."
    )
    r_tech.font.name = 'Plus Jakarta Sans'
    r_tech.font.size = Pt(10)
    r_tech.font.color.rgb = COLOR_TEXT

    # --- 5. HOW TO RUN & BUILD ---
    add_custom_heading("5. Local Development & Deployment Guide", level=1)
    
    p_run = doc.add_paragraph()
    p_run.paragraph_format.space_after = Pt(6)
    p_run.paragraph_format.line_spacing = 1.15
    r_run = p_run.add_run(
        "# Install Dependencies\n"
        "npm install\n\n"
        "# Start Development Server\n"
        "npm run dev\n\n"
        "# Build Production Static Bundle\n"
        "npm run build\n"
    )
    r_run.font.name = 'Courier New'
    r_run.font.size = Pt(9.5)
    r_run.font.color.rgb = COLOR_DARK

    # Save document
    output_path = "/Users/hum-mac/Documents/fluxrendor/Wice_Waste_Management_Documentation.docx"
    doc.save(output_path)
    print(f"Documentation DOCX generated successfully at: {output_path}")

if __name__ == "__main__":
    build_docx()
